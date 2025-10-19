import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from flask import current_app
from .s3_handler import upload_file_to_s3, get_s3_url
import requests
import io
from datetime import datetime
import re

def create_initial_book_docx(user_id):
    """
    Create an initial DOCX file for a user with dummy content
    
    Args:
        user_id (int): The user's ID
        
    Returns:
        tuple: (success (bool), message (str), s3_url (str), local_path (str))
    """
    try:
        # Create a new Document
        document = Document()
        
        # Add styles
        styles = document.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        
        # Heading style
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(18)
        heading_font.bold = True
        
        # Body style
        body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Times New Roman'
        body_font.size = Pt(12)
        
        # Add title
        title = document.add_paragraph('My Personal Book', style='CustomTitle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add author
        author = document.add_paragraph(f'User ID: {user_id}')
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add introduction
        document.add_heading('Introduction', level=1)
        intro = document.add_paragraph(
            'Welcome to your personal book! This document is yours to edit and customize. '
            'You can use the AI editing features to help you create content, or edit it manually. '
            'All changes you make will be saved to your personal copy.'
        )
        
        # Add sample chapters
        document.add_heading('Chapter 1: Getting Started', level=1)
        document.add_paragraph(
            'This is the beginning of your journey. Here you can write about your initial experiences '
            'and what you hope to achieve. The AI can help you expand on your ideas or suggest improvements.'
        )
        
        document.add_heading('Chapter 2: Developing Your Ideas', level=1)
        document.add_paragraph(
            'As you progress, you can develop your ideas further in this chapter. '
            'Think about the key concepts you want to explore and how they connect to each other.'
        )
        
        document.add_heading('Chapter 3: Putting It All Together', level=1)
        document.add_paragraph(
            'In this final chapter, you can bring all your ideas together into a cohesive whole. '
            'Consider how your various thoughts and concepts relate to each other and form a complete picture.'
        )
        
        # Create a temporary file to save the document
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, f'user_{user_id}_book.docx')
        
        # Save the document
        document.save(temp_path)
        
        # Define the S3 object name
        object_name = f'users/{user_id}/personal_book.docx'
        
        # Upload to S3
        success, message, s3_url = upload_file_to_s3(
            temp_path, 
            object_name, 
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        if not success:
            current_app.logger.error(f"Failed to upload DOCX to S3: {message}")
            return False, message, None, temp_path
        
        return True, "Personal book created successfully", s3_url, temp_path
        
    except Exception as e:
        current_app.logger.error(f"Error creating DOCX file: {str(e)}")
        return False, f"Error creating DOCX file: {str(e)}", None, None

def update_book_docx_with_ai_edits(s3_url, edit_content, edit_type='content'):
    """
    Update a DOCX file with AI-generated edits by directly modifying the document content
    
    Args:
        s3_url (str): The S3 URL of the DOCX file
        edit_content (str): The AI-generated edit content
        edit_type (str): The type of edit (content, style, structure)
        
    Returns:
        tuple: (success (bool), message (str), updated_s3_url (str))
    """
    try:
        current_app.logger.info(f"🔄 Starting AI edit process for document: {s3_url}")
        
        # Add cache-busting parameter to URL to ensure we get the latest version
        cache_bust_url = f"{s3_url}{'&' if '?' in s3_url else '?'}cache_bust={datetime.now().timestamp()}"
        current_app.logger.info(f"📥 Downloading document with cache busting: {cache_bust_url}")
        
        # Download the DOCX file from S3
        response = requests.get(cache_bust_url)
        if response.status_code != 200:
            current_app.logger.error(f"❌ Failed to download DOCX: HTTP {response.status_code}")
            return False, f"Failed to download DOCX file: HTTP {response.status_code}", None
        
        # Create a document from the downloaded content
        doc_bytes = io.BytesIO(response.content)
        document = Document(doc_bytes)
        
        # Create a temporary file for the updated document
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, f'updated_book_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx')
        
        current_app.logger.info(f"✏️ Applying AI edits to document")
        
        # Parse the edit instruction to determine what to modify
        # This is a simplified approach that looks for common edit patterns
        
        # Extract edit instruction from AI response
        edit_instruction = edit_content.lower()
        
        # Check for title change instructions
        if "title" in edit_instruction and ("change" in edit_instruction or "update" in edit_instruction or "set" in edit_instruction or "make" in edit_instruction):
            # Look for quoted text or text after "to" that might be the new title
            title_match = re.search(r'"([^"]+)"', edit_content)
            if not title_match:
                title_match = re.search(r'title\s+to\s+["\']?([^"\'.,]+)["\']?', edit_content, re.IGNORECASE)
            
            if title_match:
                new_title = title_match.group(1).strip()
                current_app.logger.info(f"🔤 Changing document title to: '{new_title}'")
                
                # Find and update the title paragraph (usually the first heading)
                title_found = False
                for i, para in enumerate(document.paragraphs):
                    if para.style.name.startswith('Heading') or i == 0:
                        para.text = new_title
                        title_found = True
                        break
                
                # If no title paragraph found, add one at the beginning
                if not title_found:
                    document.paragraphs[0].insert_paragraph_before(new_title).style = 'Heading 1'
        
        # Check for adding a section
        elif "add" in edit_instruction and ("section" in edit_instruction or "chapter" in edit_instruction or "heading" in edit_instruction):
            # Extract section title and content
            section_title = None
            section_content = None
            
            # Try to find section title
            title_match = re.search(r'section\s+["\']([^"\']+)["\']', edit_content)
            if not title_match:
                title_match = re.search(r'section\s+(?:titled|called|named)\s+["\']?([^"\'.,]+)["\']?', edit_content, re.IGNORECASE)
            
            if title_match:
                section_title = title_match.group(1).strip()
            
            # Extract content after "with content" or similar phrases
            content_match = re.search(r'with\s+(?:content|text|information)\s+["\']([^"\']+)["\']', edit_content, re.IGNORECASE)
            if content_match:
                section_content = content_match.group(1).strip()
            else:
                # Try to extract content from the remaining text
                lines = edit_content.split('\n')
                if len(lines) > 2:
                    section_content = '\n'.join(lines[2:]).strip()
            
            # Add the new section
            if section_title:
                current_app.logger.info(f"📝 Adding new section: '{section_title}'")
                document.add_heading(section_title, level=1)
                
                if section_content:
                    document.add_paragraph(section_content)
                else:
                    document.add_paragraph("Content for this section.")
        
        # For more complex edits, use a more general approach
        else:
            # Try to extract structured content from the AI response
            lines = edit_content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if this is a heading
                if line.startswith('#'):
                    # Count the number of # to determine heading level
                    heading_level = min(line.count('#', 0, line.find(' ')), 6)
                    heading_text = line.strip('#').strip()
                    document.add_heading(heading_text, level=heading_level)
                    current_paragraph = None
                
                # Check if this is a list item
                elif line.startswith('- ') or line.startswith('* '):
                    item_text = line[2:].strip()
                    document.add_paragraph(item_text, style='List Bullet')
                    current_paragraph = None
                
                # Check if this is a numbered list item
                elif re.match(r'^\d+\.', line):
                    item_text = re.sub(r'^\d+\.\s*', '', line)
                    document.add_paragraph(item_text, style='List Number')
                    current_paragraph = None
                
                # Otherwise, treat as regular paragraph text
                else:
                    if current_paragraph is None:
                        current_paragraph = document.add_paragraph(line)
                    else:
                        current_paragraph.add_run('\n' + line)
        
        # Save the updated document
        document.save(temp_path)
        current_app.logger.info(f"💾 Saved updated document to temporary file: {temp_path}")
        
        # Extract the object name from the S3 URL
        object_name = s3_url.split('/')[-2] + '/' + s3_url.split('/')[-1]
        if 'amazonaws.com' in object_name:
            # Handle full S3 URLs
            parts = s3_url.split('/')
            object_name = '/'.join(parts[3:])  # Skip https://bucket.region.amazonaws.com/
        
        # Add timestamp to object name to prevent caching issues
        if '?' in object_name:
            object_name = object_name.split('?')[0]
            
        # Upload the updated document to S3 (overwriting the existing one)
        current_app.logger.info(f"📤 Uploading updated document to S3: {object_name}")
        success, message, updated_s3_url = upload_file_to_s3(
            temp_path, 
            object_name, 
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            cache_control='no-cache, no-store, must-revalidate',
            content_disposition='inline'
        )
        
        if not success:
            current_app.logger.error(f"❌ Failed to upload updated DOCX: {message}")
            return False, f"Failed to upload updated DOCX to S3: {message}", None
        
        # Add timestamp to URL to force refresh
        if updated_s3_url:
            timestamp = int(datetime.now().timestamp())
            updated_s3_url = f"{updated_s3_url}{'&' if '?' in updated_s3_url else '?'}t={timestamp}"
            current_app.logger.info(f"✅ Document updated successfully with timestamp: {updated_s3_url}")
        
        return True, "Document updated successfully with AI edits", updated_s3_url
        
    except Exception as e:
        current_app.logger.error(f"❌ Error updating DOCX file: {str(e)}")
        return False, f"Error updating DOCX file: {str(e)}", None

def update_book_docx_with_manual_edits(s3_url, manual_content):
    """
    Update a DOCX file with manual edits
    
    Args:
        s3_url (str): The S3 URL of the DOCX file
        manual_content (str): The manually edited content
        
    Returns:
        tuple: (success (bool), message (str), updated_s3_url (str))
    """
    try:
        # Create a new document with the manual content
        document = Document()
        
        # Add the manual content paragraphs
        for paragraph in manual_content.split('\n\n'):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())
        
        # Create a temporary file for the updated document
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, 'manual_updated_book.docx')
        
        # Save the document
        document.save(temp_path)
        
        # Extract the object name from the S3 URL
        object_name = s3_url.split('/')[-2] + '/' + s3_url.split('/')[-1]
        if 'amazonaws.com' in object_name:
            # Handle full S3 URLs
            parts = s3_url.split('/')
            object_name = '/'.join(parts[3:])  # Skip https://bucket.region.amazonaws.com/
        
        # Upload the updated document to S3 (overwriting the existing one)
        success, message, updated_s3_url = upload_file_to_s3(
            temp_path, 
            object_name, 
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        if not success:
            return False, f"Failed to upload manually updated DOCX to S3: {message}", None
        
        return True, "Document updated successfully with manual edits", updated_s3_url
        
    except Exception as e:
        current_app.logger.error(f"Error updating DOCX file with manual edits: {str(e)}")
        return False, f"Error updating DOCX file with manual edits: {str(e)}", None

def extract_text_from_docx(s3_url):
    """
    Extract text content from a DOCX file
    
    Args:
        s3_url (str): The S3 URL of the DOCX file
        
    Returns:
        tuple: (success (bool), content (str))
    """
    try:
        # Download the DOCX file from S3
        response = requests.get(s3_url)
        if response.status_code != 200:
            return False, f"Failed to download DOCX file: HTTP {response.status_code}"
        
        # Create a document from the downloaded content
        doc_bytes = io.BytesIO(response.content)
        document = Document(doc_bytes)
        
        # Extract text from paragraphs
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        
        return True, '\n\n'.join(full_text)
        
    except Exception as e:
        current_app.logger.error(f"Error extracting text from DOCX: {str(e)}")
        return False, f"Error extracting text from DOCX: {str(e)}" 