"""
Text Extraction Service
Handles text extraction from various document types: PDF, DOCX, TXT, Images (AWS Bedrock Vision)
"""
import logging
from pathlib import Path
from typing import Optional, Tuple
import tempfile
import os
import base64
import json
import boto3
from config.settings import settings

# PDF extraction
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# DOCX extraction
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Image handling
try:
    from PIL import Image
    HAS_IMAGE = True
except ImportError:
    HAS_IMAGE = False

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Service for extracting text from various document types"""
    
    def __init__(self):
        """Initialize Bedrock client for image analysis"""
        self.bedrock_client = boto3.client(
            'bedrock-runtime',
            region_name=settings.AWS_REGION,
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY
        )
        self.vision_model_id = settings.BEDROCK_MODEL_ID
    
    async def extract_text_from_file(
        self,
        file_path: str,
        file_type: str,
        mime_type: Optional[str] = None,
        user_dog_profiles: Optional[list] = None
    ) -> Tuple[str, Optional[str]]:
        """
        Extract text from a file
        
        Args:
            file_path: Path to the file
            file_type: File extension (pdf, docx, txt, jpg, png, etc.)
            mime_type: MIME type of the file
            user_dog_profiles: List of user's dog profiles for personalized image analysis
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            file_type = file_type.lower().lstrip('.')
            
            # Text files
            if file_type in ['txt', 'text']:
                return await TextExtractionService._extract_from_text(file_path)
            
            # PDF files
            elif file_type == 'pdf':
                if not HAS_PDF:
                    return "", "PyPDF2 not installed - cannot extract PDF"
                return await TextExtractionService._extract_from_pdf(file_path)
            
            # DOCX files
            elif file_type in ['docx', 'doc']:
                if not HAS_DOCX:
                    return "", "python-docx not installed - cannot extract DOCX"
                return await TextExtractionService._extract_from_docx(file_path)
            
            # Image files (AWS Bedrock Vision)
            elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif', 'webp']:
                if not HAS_IMAGE:
                    return "", "PIL not installed - cannot handle images"
                return await self._extract_from_image(file_path, user_dog_profiles)
            
            else:
                return "", f"Unsupported file type: {file_type}"
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            return "", str(e)
    
    @staticmethod
    async def _extract_from_text(file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text.strip(), None
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return text.strip(), None
            except Exception as e:
                return "", f"Failed to read text file: {str(e)}"
        except Exception as e:
            return "", f"Failed to read text file: {str(e)}"
    
    @staticmethod
    async def _extract_from_pdf(file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from PDF file"""
        try:
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            full_text = "\n\n".join(text_parts)
            return full_text.strip(), None
            
        except Exception as e:
            return "", f"Failed to extract PDF: {str(e)}"
    
    @staticmethod
    async def _extract_from_docx(file_path: str) -> Tuple[str, Optional[str]]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            return full_text.strip(), None
            
        except Exception as e:
            return "", f"Failed to extract DOCX: {str(e)}"
    
    async def _extract_from_image(self, file_path: str, user_dog_profiles: Optional[list] = None) -> Tuple[str, Optional[str]]:
        """Extract text and describe image using personalized vision service"""
        try:
            if not HAS_IMAGE:
                return "", "PIL not installed for image handling"
            
            # Read and encode image to base64
            with open(file_path, 'rb') as img_file:
                image_data = base64.b64encode(img_file.read()).decode('utf-8')
            
            # Import vision service here to avoid circular imports
            from services.vision_service import vision_service
            
            # Get filename for context
            filename = Path(file_path).name
            
            # Use personalized vision service if dog profiles are available
            if user_dog_profiles:
                logger.info(f"🖼️ Analyzing image with personalized dog context...")
                description = await vision_service.analyze_chat_image(
                    image_data=image_data,
                    user_dog_profiles=user_dog_profiles,
                    filename=filename
                )
            else:
                logger.info(f"🖼️ Analyzing image with generic description...")
                description = await vision_service._generate_generic_image_description(
                    image_data=image_data,
                    filename=filename
                )
            
            logger.info(f"✅ Image analysis complete: {len(description)} chars")
            return description.strip(), None
            
        except Exception as e:
            logger.error(f"❌ Image analysis failed: {str(e)}", exc_info=True)
            return "", f"Failed to analyze image: {str(e)}"
    
    @staticmethod
    def get_supported_file_types() -> dict:
        """Get list of supported file types and their status"""
        return {
            "text": {"supported": True, "extensions": [".txt"]},
            "pdf": {"supported": HAS_PDF, "extensions": [".pdf"]},
            "docx": {"supported": HAS_DOCX, "extensions": [".docx", ".doc"]},
            "images": {
                "supported": HAS_IMAGE,  # Using AWS Bedrock Vision
                "extensions": [".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".gif", ".webp"]
            }
        }
